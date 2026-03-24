import os
from docx import Document

doc_path = os.path.join(os.path.dirname(__file__), '../Sample_DRs/17.02.2026 IT DEPT  format for Daily Report.docx')
doc = Document(doc_path)

print("--- Table 0 ---")
table_0 = doc.tables[0]
for r, row in enumerate(table_0.rows):
    print(f"Row {r}: {[cell.text.strip() for cell in row.cells]}")

print("\n--- Paragraphs containing Keywords ---")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Placement" in text or "Any other matter" in text or "MTP" in text or "matter" in text.lower() or "placement" in text.lower():
        print(f"[{i}] {text}")
        # Print next 5 paragraphs
        for j in range(1, 6):
            if i + j < len(doc.paragraphs):
                # Only print if it's not empty
                next_text = doc.paragraphs[i + j].text.strip()
                if next_text:
                    print(f"  [{i+j}] {next_text}")
        print("-------")
