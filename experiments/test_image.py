import os
from docx import Document
from docx.shared import Inches

doc_path = os.path.join(os.path.dirname(__file__), '../Sample_DRs/17.02.2026 Daily Report.docx')
doc = Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
images_found = 0

for i, p in enumerate(doc.paragraphs):
    for r_idx, run in enumerate(p.runs):
        if 'Graphic' in run._element.xml or 'pic:pic' in run._element.xml or 'drawing' in run._element.xml:
            print(f"Found image in paragraph {i}, run {r_idx}")
            images_found += 1
            
print(f"Images found: {images_found}")

# Alternatively, check doc.inline_shapes
print(f"Inline shapes count: {len(doc.inline_shapes)}")
