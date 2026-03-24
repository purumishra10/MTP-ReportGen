import os
from docx import Document

doc_path = os.path.join(os.path.dirname(__file__), '../Sample_DRs/Empty Daily DR.docx')
doc = Document(doc_path)

print("--- Master Template Table 2 ---")
table_2 = doc.tables[2]
for r, row in enumerate(table_2.rows):
    print(f"  Row {r}: {[cell.text.strip() for cell in row.cells]}")

print("\n\n--- IT Dept  Document ---")
dept_doc_path = os.path.join(os.path.dirname(__file__), '../Sample_DRs/17.02.2026 IT DEPT  format for Daily Report.docx')
dept_doc = Document(dept_doc_path)

print("\n--- IT Dept Tables ---")
for i, table in enumerate(dept_doc.tables):
    if len(table.rows) > 0:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"Table {i} first row: {first_row}")
        
print("\n--- IT Dept Paragraphs (first 100) ---")
for i, p in enumerate(dept_doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"[{i}] {text[:50]}")
    if i > 50:
        break
