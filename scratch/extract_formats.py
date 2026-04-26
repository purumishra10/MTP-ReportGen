import docx
import os
import json

def extract_structure(file_path):
    doc = docx.Document(file_path)
    structure = []
    
    current_section = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.endswith(':') or text.isupper() or any(text.startswith(prefix) for prefix in ['I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.']):
            current_section = text
            structure.append({"section": text, "type": "paragraph"})
            
    for table in doc.tables:
        header = []
        if len(table.rows) > 0:
            header = [cell.text.strip() for cell in table.rows[0].cells]
        structure.append({"type": "table", "header": header})
        
    return structure

files = [
    r"c:\Users\Shiva\MTP-ReportGen\March_2026\30th March 2026\Library Daily Report.docx",
    r"c:\Users\Shiva\MTP-ReportGen\March_2026\30th March 2026\Staff & Student attendance report-30-03-26.docx",
    r"c:\Users\Shiva\MTP-ReportGen\March_2026\31st March 2026\Daily Report 31st March 2026-MTP.docx"
]

results = {}
for f in files:
    if os.path.exists(f):
        results[os.path.basename(f)] = extract_structure(f)
    else:
        results[os.path.basename(f)] = "File not found"

print(json.dumps(results, indent=2))
