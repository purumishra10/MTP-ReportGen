import os
import glob
from docx import Document
import re
import json

base_path = r"C:\Users\Shiva\MTP-ReportGen\backend\data\March_2026\March_2026"
folders = ["24th March 2026", "25th March 2026", "26th March 2026", "30th March 2026", "31st March 2026"]

def normalize_dept_name(filename):
    name = os.path.basename(filename)
    name = name.replace(".docx", "")
    # Remove dates and 'Daily Report' text from start or end
    name = re.sub(r'(?i)^\s*daily report\s*-?\s*', '', name)
    name = re.sub(r'(?i)\s*-?\s*daily report.*$', '', name)
    name = re.sub(r'(?i)format for daily report.*$', '', name)
    name = re.sub(r'(?i)^\s*staff\s*&?\s*student attendance report.*$', 'Staff Student Attendance', name)
    name = re.sub(r'(?i)^\s*staff attendance report.*$', 'Staff Attendance', name)
    
    # Strip dates like -25.03.2026 or 25-03-2026 or 25th March 2026
    name = re.sub(r'\d{1,2}(st|nd|rd|th)?\s*[a-zA-Z]+\s*\d{4}.*$', '', name)
    name = re.sub(r'[-@]*\d{1,2}[.-]\d{1,2}[.-]\d{2,4}.*$', '', name)
    
    name = name.strip(" -_-@")
    if not name:
        name = "Unknown"
    return name

def extract_format_from_docx(filepath):
    try:
        doc = Document(filepath)
    except Exception as e:
        if "CRC-32" in str(e):
            print(f"Fixing CRC error for {filepath} by stripping media...")
            import zipfile
            import io
            temp_path = filepath + ".tmp.docx"
            try:
                zin = zipfile.ZipFile(filepath, 'r')
                zout = zipfile.ZipFile(temp_path, 'w')
                for item in zin.infolist():
                    if not item.filename.startswith('word/media/'):
                        buffer = zin.read(item.filename)
                        zout.writestr(item, buffer)
                zin.close()
                zout.close()
                doc = Document(temp_path)
                os.remove(temp_path)
            except Exception as e2:
                print(f"Recovery failed for {filepath}: {e2}")
                return []
        else:
            print(f"Error reading {filepath}: {e}")
            return []
        
    tables_format = []
    
    # Simple extraction: just look at all tables in order. 
    # Try to find a preceding paragraph that looks like a title.
    body_elements = doc.element.body
    
    current_title = "Unknown Section"
    for elem in body_elements:
        if elem.tag.endswith('p'):
            text = elem.text.strip()
            if text and len(text) < 100:
                # possible title
                current_title = text
        elif elem.tag.endswith('tbl'):
            # Find the corresponding table object
            for t in doc.tables:
                if t._element == elem:
                    if len(t.rows) > 0:
                        # Extract headers from first row
                        headers = []
                        for cell in t.rows[0].cells:
                            h = cell.text.strip()
                            h = re.sub(r'\s+', ' ', h)
                            if h not in headers: # prevent duplicates from merged cells
                                headers.append(h)
                        tables_format.append({
                            "section_title": current_title,
                            "columns": headers
                        })
                    break
    return tables_format

# Extract Formats per Day
dept_formats = {}
file_sources = {}

for folder in folders:
    folder_path = os.path.join(base_path, folder)
    if not os.path.exists(folder_path):
        continue
    files = glob.glob(os.path.join(folder_path, "*.docx"))
    
    for file in files:
        dept = normalize_dept_name(file)
        if dept not in dept_formats:
            dept_formats[dept] = {}
        if dept not in file_sources:
            file_sources[dept] = {}
            
        fmt = extract_format_from_docx(file)
        dept_formats[dept][folder] = fmt
        file_sources[dept][folder] = os.path.basename(file)

# Compute Intersections & Discarded
final_formats = {}
discarded_info = []

for dept, days_data in dept_formats.items():
    # Only process if they have reports on multiple days or all available days.
    # We will intersect available days.
    available_days = list(days_data.keys())
    if not available_days:
        continue
        
    # Start intersection with day 1
    base_day = available_days[0]
    intersecting_format = []
    
    # We will match tables by columns to form intersection
    for table_template in days_data[base_day]:
        # Check if a functionally similar table exists in all other days
        columns_set_base = set(table_template['columns'])
        
        present_in_all = True
        for day in available_days[1:]:
            found = False
            for target_table in days_data[day]:
                target_col_set = set(target_table['columns'])
                # If they share at least 50% columns, we consider it same table layout
                if len(columns_set_base.intersection(target_col_set)) >= len(columns_set_base) * 0.5:
                    found = True
                    break
            if not found:
                present_in_all = False
                break
                
        if present_in_all:
            intersecting_format.append(table_template)
            
    final_formats[dept] = intersecting_format

    # Record discarded
    for day in available_days:
        for table_template in days_data[day]:
            # is it in intersecting?
            columns_set_base = set(table_template['columns'])
            found_in_intersect = False
            for inter_t in intersecting_format:
                inter_col_set = set(inter_t['columns'])
                if len(columns_set_base.intersection(inter_col_set)) >= len(columns_set_base) * 0.5:
                    found_in_intersect = True
                    break
            if not found_in_intersect:
                discarded_info.append({
                    "department": dept,
                    "date": day,
                    "file": file_sources[dept][day],
                    "section_title": table_template['section_title'],
                    "columns": table_template['columns']
                })

# Grouping Similar Departments
grouped_formats = []
for dept, fmt in final_formats.items():
    # Attempt to group with existing
    placed = False
    for group in grouped_formats:
        group_fmt = group['format']
        # Compare group format with this fmt
        if len(group_fmt) == len(fmt):
            # Check if tables match roughly (allow 1-2 missing columns or extra)
            matches = 0
            for t1, t2 in zip(group_fmt, fmt):
                set1 = set(t1['columns'])
                set2 = set(t2['columns'])
                diff = len(set1.symmetric_difference(set2))
                if diff <= 2:
                    matches += 1
            if matches == len(fmt):
                group['departments'].append(dept)
                placed = True
                break
    if not placed:
        grouped_formats.append({
            "departments": [dept],
            "format": fmt
        })

# Create Directories
json_dir = r"C:\Users\Shiva\MTP-ReportGen\backend\formats\json"
md_dir = r"C:\Users\Shiva\MTP-ReportGen\backend\formats\human_readable"
os.makedirs(json_dir, exist_ok=True)
os.makedirs(md_dir, exist_ok=True)

# Save JSON and MD
for idx, group in enumerate(grouped_formats):
    depts = group['departments']
    group_name = depts[0]
    if len(depts) > 1:
        if "ME" in depts or "CSE" in depts:
             group_name = "Engineering_Depts"
        elif "Chemistry" in depts or "English" in depts:
             group_name = "Subject_Depts"
        else:
             group_name = f"Group_{idx+1}"
    
    # Instead of group_name we will just use the first department name as base if its a group
    # But user said: "use the same name as the name of depts as used in file names"
    # we can name the file like "ME_Civil_CSE.json" or pick one. 
    safe_name = "_".join(depts)[:100] # truncate if too long
    safe_name = re.sub(r'[^a-zA-Z0-9_]', '_', safe_name)
    
    json_path = os.path.join(json_dir, f"{safe_name}.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(group, f, indent=4)
        
    md_path = os.path.join(md_dir, f"{safe_name}.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# Format Group: {', '.join(depts)}\n\n")
        for table in group['format']:
            f.write(f"## {table['section_title']}\n")
            f.write("| " + " | ".join(table['columns']) + " |\n")
            f.write("|" + "|".join(["---"] * len(table['columns'])) + "|\n")
            f.write("\n")

# Save Discarded
discarded_path = os.path.join(md_dir, "discarded_fields.md")
with open(discarded_path, 'w', encoding='utf-8') as f:
    f.write("# Discarded / Extra Fields (Bullshit Data)\n\n")
    for info in discarded_info:
        f.write(f"**Department:** {info['department']}\n")
        f.write(f"**Date:** {info['date']}\n")
        f.write(f"**File:** {info['file']}\n")
        f.write(f"**Section Title:** {info['section_title']}\n")
        f.write(f"**Columns Found:** {', '.join(info['columns'])}\n")
        f.write("---\n")

print("Format extraction completed successfully.")
