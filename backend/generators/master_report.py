import os
import io
from docx import Document
from docx.shared import Inches

def generate_master_report(template_path, output_path, attendance_data, mtp_data, image_data, missing_depts, library_data, executive_summary=None):
    doc = Document(template_path)

    try:
        table_2 = doc.tables[2]
        att_cell = table_2.rows[0].cells[1]
        att_cell.text = "Staff & Student Attendance Report : Attached.\n\nConsolidated Absent Faculty:\n"
        
        if not attendance_data and not missing_depts:
            att_cell.text += "Data Not Received\n"
        else:
            for entry in attendance_data:
                att_cell.text += f"- {entry}\n"
            
        if missing_depts:
            att_cell.text += "\nMissing Reports From:\n"
            for dept in missing_depts:
                att_cell.text += f"- {dept} Data Not Received\n"

        mtp_cell = table_2.rows[1].cells[1]
        mtp_cell.text = "MTP:\n\n"
        
        if executive_summary:
            p = mtp_cell.add_paragraph()
            r = p.add_run("INSTITUTIONAL EXECUTIVE SUMMARY:\n")
            r.bold = True
            r.underline = True
            mtp_cell.add_paragraph(executive_summary)
            mtp_cell.add_paragraph("\n")
            
        if not mtp_data and not missing_depts:
            mtp_cell.text += "Data Not Received\n"
            
        placed_depts = set()
        for row in table_2.rows:
            if len(row.cells) > 1:
                cell = row.cells[1]
                text = cell.text.strip()
                if text.startswith("HOD-"):
                    dept_key = text.replace("HOD-", "").replace(":", "").strip()
                    if dept_key in mtp_data:
                        cell.text = f"HOD-{dept_key}:\n"
                        for p_text in mtp_data[dept_key]:
                            if p_text:
                                cell.add_paragraph(p_text)
                        placed_depts.add(dept_key)
                        
        for dept, paragraphs in mtp_data.items():
            if dept not in placed_depts:
                new_row = table_2.add_row()
                if len(new_row.cells) > 1:
                    new_row.cells[0].text = ""  # blank for numeral column
                    new_cell = new_row.cells[1]
                    new_cell.text = f"HOD-{dept}:\n"
                    for p_text in paragraphs:
                        if p_text:
                            new_cell.add_paragraph(p_text)

        for img_bytes, ext in image_data:
            try:
                p = mtp_cell.add_paragraph()
                p.alignment = 1 
                r = p.add_run()
                img_stream = io.BytesIO(img_bytes)
                r.add_picture(img_stream, width=Inches(4.5))
            except Exception as e:
                print(f"Failed to insert image - {e}")
                
        # Library data
        if library_data and len(doc.tables) > 3:
            table_3 = doc.tables[3]
            try:
                if len(table_3.rows) > 1:
                    table_3.rows[1].cells[2].text = str(library_data.get("books_issued", ""))
                if len(table_3.rows) > 2:
                    table_3.rows[2].cells[2].text = str(library_data.get("books_returned", ""))
                if len(table_3.rows) > 3:
                    table_3.rows[3].cells[2].text = str(library_data.get("visits", ""))
            except Exception as e:
                print(f"Error populating library table: {e}")
                
    except Exception as e:
        print(f"Error populating document: {e}")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
