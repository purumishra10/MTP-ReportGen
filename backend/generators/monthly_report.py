import os
from docx import Document
from database import get_monthly_records
from datetime import datetime

def generate_monthly_report(output_dir):
    os.makedirs(output_dir, exist_ok=True)
    month_name = datetime.now().strftime("%B_%Y")
    output_path = os.path.join(output_dir, f"MTP_Monthly_Report_{month_name}.docx")
    
    doc = Document()
    doc.add_heading(f'MTP Monthly Consolidated Report - {month_name.replace("_", " ")}', 0)
    
    records = get_monthly_records()
    
    if not records:
        doc.add_paragraph("No MTP records found in the database.")
    else:
        current_date = None
        for r_date, dept, text in records:
            if current_date != r_date:
                current_date = r_date
                doc.add_heading(f"Date: {current_date}", level=1)
                
            doc.add_heading(f"Department: {dept}", level=2)
            for p in text.split('\n'):
                if p.strip():
                    doc.add_paragraph(p.strip())
            
    doc.save(output_path)
    return output_path
