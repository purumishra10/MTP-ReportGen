"""
Deterministic Structured Extractor (v3)
========================================
Reads DOCX tables directly and extracts structured data WITHOUT any LLM.
Uses header-matching (not position-based indexing) to handle varying table orders.

v3 improvements over v2:
- DOCX repair: skips corrupt media (BadZipFile) so extraction still works
- _get_header_text() scans ALL rows (not just first 2) and normalises embedded \\n
- Broader participation fingerprint — handles variant header wordings
- MTP nested table detection scans ALL rows, not just header rows
- Empty events table detected → full table text still forwarded so LLM can read
  context, PLUS loose_paragraphs always forwarded
- Infrastructure fingerprint handles "Reported\\nOn" (newline in header cell)
- Loose paragraphs labelled with dept context before LLM forwarding
- Attendance numbers are 100% deterministic — never touched by LLM
"""

import io
import os
import re
import zipfile
from typing import Optional

from docx import Document


# ── Header fingerprints for table identification ──────────────────────────────
# Each fingerprint is a set of keywords that MUST appear in the header area.
# We normalize headers to lowercase before matching.

FINGERPRINTS = {
    "attendance": {"on rolls", "absent"},
    "infrastructure": {"description", "reported"},
    "events": {"event", "duration"},
    "staff_participation": {"name", "event", "status"},
    "student_participation": {"name", "event", "status"},
    "staff_changes": {"name", "designation", "joining"},
    "classwork": {"dept", "subject", "adjusted"},
    "incidents": {"name", "brief", "remarks"},
    # Library-specific
    "library_transactions": {"particulars", "no"},
    "library_attendance": {"on roll", "absent", "present"},
    # Overall Attendance report specific
    "overall_staff_attendance": {"on rolls", "overall performance"},
    "overall_student_attendance": {"ce", "eee", "me", "ece", "cse"},
    # MTP specific
    "mtp_narrative": {"mtp:"},
    "mtp_batch_pills": {"batch pills"},
}

# Sections whose data should be extracted deterministically (no LLM)
DETERMINISTIC_SECTIONS = {
    "attendance", "infrastructure", "staff_changes",
    "classwork", "incidents", "library_transactions", "library_attendance",
    "overall_staff_attendance", "overall_student_attendance",
}

# Sections whose data goes to LLM for summarization
NARRATIVE_SECTIONS = {
    "events", "staff_participation", "student_participation",
    "mtp_narrative", "mtp_batch_pills",
}


# ── DOCX repair helper ────────────────────────────────────────────────────────

def _repair_docx(file_bytes: bytes) -> bytes:
    """
    Strip corrupt media files from a DOCX (which is a ZIP) and patch .rels XML
    so python-docx doesn't try to reference the now-missing files.
    """
    skipped = set()

    # Pass 1: identify corrupt entries
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zin:
            for item in zin.infolist():
                try:
                    zin.read(item.filename)
                except Exception:
                    skipped.add(item.filename)
                    print(f"[REPAIR] Found corrupt entry: {item.filename}")
    except Exception as e:
        print(f"[REPAIR] Could not scan ZIP: {e}")
        return file_bytes

    if not skipped:
        return file_bytes  # Nothing to fix

    # Collect basenames for rels matching (e.g. "image1.png" from "word/media/image1.png")
    skipped_basenames = {s.split("/")[-1] for s in skipped}

    # Pass 2: rebuild ZIP without corrupt entries, patching .rels files inline
    buf = io.BytesIO()
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zin, \
             zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in skipped:
                    continue  # Drop corrupt file entirely
                try:
                    data = zin.read(item.filename)
                    if item.filename.endswith(".rels"):
                        data = _patch_rels_xml(data, skipped_basenames)
                    zout.writestr(item, data)
                except Exception as write_err:
                    print(f"[REPAIR] Skipping unreadable entry {item.filename}: {write_err}")
    except Exception as e:
        print(f"[REPAIR] Could not repair DOCX: {e}")
        return file_bytes

    buf.seek(0)
    return buf.read()


def _patch_rels_xml(rels_data: bytes, skipped_basenames: set) -> bytes:
    """
    Remove <Relationship .../> elements from a .rels XML file that point to
    any of the skipped (corrupt/missing) files, identified by basename.
    Uses regex so we don't need to parse XML (which may itself be malformed).
    """
    import re as _re
    try:
        text = rels_data.decode("utf-8", errors="replace")
        for basename in skipped_basenames:
            escaped = _re.escape(basename)
            # Target="media/image1.png" or Target="../media/image1.png"
            # The basename appears after a '/' or at the start of the Target value
            text = _re.sub(
                r'<Relationship\b[^>]*Target=["\'][^"\']*(?:/|^)' + escaped + r'["\'][^/]*/?>',
                "",
                text,
            )
        return text.encode("utf-8")
    except Exception:
        return rels_data  # If patching fails, return original



def extract_structured_data(file_bytes: bytes, dept_code: str,
                            dept_name: str, is_library: bool = False) -> dict:
    """
    Extract structured data from a DOCX file.

    Returns a dict with:
    - Deterministically extracted sections (attendance, infra, etc.)
    - Raw text for narrative sections (events, participation)
    - Loose paragraphs as raw text (always forwarded to LLM)

    IMPORTANT: Attendance numbers are NEVER sent to the LLM.
    They are always extracted deterministically from the table.
    """
    # ── Open DOCX, repairing corrupt media if necessary ──────────────────────
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        err_str = str(e).lower()
        if "crc" in err_str or "zip" in err_str or "bad" in err_str or "no item" in err_str:
            print(f"[WARN] Corrupt DOCX for {dept_code}, attempting repair: {e}")
            repaired = _repair_docx(file_bytes)
            doc = Document(io.BytesIO(repaired))
        else:
            raise

    tables = doc.tables

    # Classify each table by matching its header area
    classified = _classify_tables(tables, is_library)

    result = {
        "dept_code": dept_code,
        "dept_name": dept_name,
        # Deterministic extractions — NEVER sent to LLM
        "attendance": None,
        "infrastructure_issues": [],
        "staff_changes": [],
        "classwork_adjustment_count": 0,
        "incidents": [],
        # Library-specific
        "library_attendance": None,
        "library_transactions": {},
        "library_services": {},
        # Overall Attendance (from separate doc)
        "overall_staff_attendance_table": [],
        "overall_student_attendance_table": [],
        # MTP specific (sent to LLM as narrative)
        "mtp_narrative": "",
        "mtp_batch_pills": "",
        # Narrative text for LLM — all of these go to ai_service.py
        "events_text": "",
        "staff_participation_text": "",
        "student_participation_text": "",
        "other_matters_text": "",
        "loose_paragraphs": "",
    }

    # Track which participation table we've seen (first = staff, second = student)
    participation_count = 0

    for section_type, table in classified:
        if section_type == "attendance" and not is_library:
            # DETERMINISTIC — numbers never go to LLM
            result["attendance"] = _extract_attendance(table, dept_code, dept_name)

        elif section_type == "library_attendance" and is_library:
            # DETERMINISTIC
            result["library_attendance"] = _extract_library_attendance(table)

        elif section_type == "infrastructure":
            # DETERMINISTIC
            result["infrastructure_issues"] = _extract_infrastructure(table, dept_name)

        elif section_type == "staff_changes":
            # DETERMINISTIC
            result["staff_changes"] = _extract_staff_changes(table, dept_name)

        elif section_type == "classwork":
            # DETERMINISTIC — just count non-empty rows
            result["classwork_adjustment_count"] = _count_non_empty_rows(table)

        elif section_type == "incidents":
            # DETERMINISTIC
            result["incidents"] = _extract_incidents(table, dept_name)

        elif section_type == "library_transactions":
            # DETERMINISTIC
            txn, svc = _extract_library_data(table)
            result["library_transactions"] = txn
            result["library_services"] = svc

        elif section_type == "events":
            # Forward full table text to LLM regardless of whether rows look empty.
            # Many depts leave the table blank but write events as loose paragraphs.
            table_text = _table_to_text("Events / Seminars / Workshops", table)
            if table_text.strip():
                result["events_text"] = table_text

        elif section_type in ("staff_participation", "student_participation"):
            # Disambiguate: first occurrence = staff, second = student
            if participation_count == 0:
                result["staff_participation_text"] = _table_to_text(
                    "Participation by Staff", table
                )
            else:
                result["student_participation_text"] = _table_to_text(
                    "Participation by Students", table
                )
            participation_count += 1

        elif section_type == "overall_staff_attendance":
            # DETERMINISTIC — read raw rows for report_generator
            result["overall_staff_attendance_table"] = _read_table_rows(table)

        elif section_type == "overall_student_attendance":
            # DETERMINISTIC
            result["overall_student_attendance_table"] = _read_table_rows(table)

        elif section_type == "mtp_section_iv":
            # Nested MTP table — extract narrative from nested tables
            mtp_text, pills_text = _extract_mtp_section_iv(table)
            if mtp_text:
                result["mtp_narrative"] = mtp_text
            if pills_text:
                result["mtp_batch_pills"] = pills_text

        elif section_type == "mtp_narrative":
            # Fallback: flat table containing MTP: header
            if not result["mtp_narrative"]:
                result["mtp_narrative"] = _table_to_text("MTP Section IV", table)

        elif section_type == "mtp_batch_pills":
            if not result["mtp_batch_pills"]:
                result["mtp_batch_pills"] = _table_to_text(
                    "Batch Pills Open Summary", table
                )

        # Unclassified tables go to "other matters" for LLM
        elif section_type == "unknown":
            text = _table_to_text("Other", table)
            if text.strip() and _has_real_content(text):
                result["other_matters_text"] += text + "\n\n"

    # ── Loose paragraphs (free text outside tables) ───────────────────────────
    # ONLY collect actual substantive paragraphs typed by departments.
    # Exclude standard Word template field labels and section prompts.
    TEMPLATE_PROMPTS = (
        "daily report", "staff attendance", "students attendance", "student attendance",
        "infrastructure issues", "event (short term course", "event organized",
        "participation by staff", "participation by students",
        "staff joined or left", "classwork adjustment", "incident if any",
        "any other matter to be reported", "vnr vignana jyothi",
        "department:", "b. tech", "b.tech", "m. tech", "m.tech", "minor degree"
    )

    loose = []
    for para in doc.paragraphs:
        # Normalize non-breaking spaces
        text = para.text.replace("\xa0", " ").strip()
        if not text or len(text) < 15:
            continue
        low = text.lower()
        # Skip standard template labels
        if any(low.startswith(p) or p in low and len(low) < 55 for p in TEMPLATE_PROMPTS):
            continue
        # Skip pure department name repeats
        if low == dept_name.lower() or low in ("mathematics and management sciences", "english", "civil engineering"):
            continue
        loose.append(text)

    if loose:
        result["loose_paragraphs"] = (
            f"[Free Text — {dept_name}]\n" + "\n".join(loose)
        )

    return result


# ── MTP Section IV — Nested table extraction ─────────────────────────────────

# Branch codes that appear in the Batch Pills header row
_BATCH_PILLS_BRANCH_CODES = {"CSE", "CSBS", "AIML", "IOT", "DS", "CyS", "AIDS", "IT", "ECE", "EEE", "EIE", "ME", "CE", "AE"}


def _is_batch_pills_header(row_text: str) -> bool:
    """Return True if the row looks like the branch-code header for batch pills."""
    cols = [c.strip() for c in row_text.split("|")]
    matched = sum(1 for c in cols if c in _BATCH_PILLS_BRANCH_CODES)
    return matched >= 5  # At least 5 branch codes present


def _is_batch_pills_values(row_text: str) -> bool:
    """Return True if the row is all-numeric values (the pills count row)."""
    cols = [c.strip() for c in row_text.split("|")]
    numeric_cols = [c for c in cols if c.isdigit()]
    return len(numeric_cols) >= 5


def _is_hod_row(row_text: str) -> bool:
    """Return True if the row is a HOD/section label row (not narrative content)."""
    low = row_text.lower()
    # Rows like "V | HOD-CSE (CSDS&AIDS):" or "VI | HOD-IT:"
    return ("hod" in low and "|" in row_text) or ("|" in row_text and re.match(r'^\s*(I{1,3}|IV|V{0,3}I{0,3})\s*\|', row_text))


def _extract_mtp_section_iv(table) -> tuple[str, str]:
    """
    Extract MTP Section IV data from a table that contains nested tables.

    The MTP report has a structure like:
    Row 0: "III. Students Attendance"
    Row 1: "IV" | "MTP:" + [nested table with actual narrative + batch pills]

    Returns: (mtp_narrative_text, batch_pills_text)
    """
    mtp_narrative_lines: list[str] = []
    pills_header: str | None = None
    pills_values: str | None = None
    pills_header_text: str = ""  # Title line like "2025-2026 Batch Pills Open Summary"

    def _collect_from_rows(nested_rows: list[list[str]]):
        nonlocal pills_header, pills_values, pills_header_text
        for nr in nested_rows:
            row_text = " | ".join(c for c in nr if c).strip()
            if not row_text or len(row_text) < 4:
                continue

            # Skip HOD/roman-numeral label rows (noise)
            if _is_hod_row(row_text):
                continue

            # Detect the branch-code header row for Batch Pills
            if _is_batch_pills_header(row_text):
                pills_header = row_text
                continue

            # Detect the numeric values row (must follow after header)
            if pills_header is not None and pills_values is None and _is_batch_pills_values(row_text):
                pills_values = row_text
                continue

            # Detect title line for batch pills section
            row_lower = row_text.lower()
            if ("batch" in row_lower and ("pli" in row_lower or "pills" in row_lower)
                    and ("summary" in row_lower or "open" in row_lower)):
                pills_header_text = row_text
                continue

            # Otherwise it's MTP narrative — skip if we've already started pills
            if pills_header is None:
                # Not yet in pills section — narrative
                if row_text and len(row_text) > 5:
                    mtp_narrative_lines.append(row_text)

    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip().lower()

            # Check if this cell contains nested tables (the MTP data cell)
            if cell.tables:
                for nested_table in cell.tables:
                    nested_rows = _read_table_rows(nested_table)
                    _collect_from_rows(nested_rows)

                    # Also check nested-nested tables (deep-embedded pills table)
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            if ncell.tables:
                                for deep_table in ncell.tables:
                                    deep_rows = _read_table_rows(deep_table)
                                    _collect_from_rows(deep_rows)

            # Non-table cell with MTP narrative ("MTP:" label cells)
            elif "mtp:" in cell_text or "mtp (" in cell_text:
                cell_paragraphs = []
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t and t.lower() != "mtp:" and len(t) > 3:
                        cell_paragraphs.append(t)
                if cell_paragraphs:
                    mtp_narrative_lines.extend(cell_paragraphs)

    # Build clean batch pills text: title + pipe-table only
    batch_pills = ""
    if pills_header_text:
        batch_pills = pills_header_text + "\n"
    if pills_header and pills_values:
        batch_pills += pills_header + "\n" + pills_values
    elif pills_header:
        batch_pills += pills_header

    return "\n".join(mtp_narrative_lines).strip(), batch_pills.strip()


# ── Table classification ─────────────────────────────────────────────────────

def _classify_tables(tables, is_library: bool) -> list[tuple[str, object]]:
    """Classify each table by matching its full header area against fingerprints."""
    classified = []
    participation_seen = 0

    for table in tables:
        # Scan ALL rows (not just top 2) to catch MTP and non-standard layouts
        header_text = _get_header_text(table, max_rows=None)
        section_type = _match_fingerprint(header_text, is_library)

        # Special: detect MTP Section IV table (has nested tables anywhere)
        if section_type in ("mtp_narrative", "unknown"):
            if _table_has_nested_tables(table):
                # Full-table scan for "mtp:" keyword
                if "mtp:" in header_text or "mtp (" in header_text:
                    section_type = "mtp_section_iv"
                elif section_type == "unknown":
                    # Could still be an MTP section — check all cell text
                    all_cell_text = _get_all_cell_text(table)
                    if "mtp:" in all_cell_text.lower():
                        section_type = "mtp_section_iv"

        # Handle dual participation tables (same headers, different meaning)
        if section_type == "staff_participation":
            if participation_seen == 0:
                section_type = "staff_participation"
            else:
                section_type = "student_participation"
            participation_seen += 1

        classified.append((section_type, table))

    return classified


def _table_has_nested_tables(table) -> bool:
    """Check if any cell in the table contains nested tables."""
    for row in table.rows:
        for cell in row.cells:
            if cell.tables:
                return True
    return False


def _get_all_cell_text(table) -> str:
    """Get all text from all cells in a table (for deep keyword search)."""
    texts = []
    for row in table.rows:
        seen = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            t = cell.text.strip().lower()
            if t:
                texts.append(t)
    return " ".join(texts)


def _get_header_text(table, max_rows: Optional[int] = None) -> str:
    """
    Get all text from table rows, normalized.

    v3 change: scans ALL rows by default (max_rows=None), normalises
    embedded newlines so 'Reported\\nOn' becomes 'Reported On' and
    correctly matches the 'reported' fingerprint keyword.
    """
    texts = []
    rows = table.rows if max_rows is None else list(table.rows)[:max_rows]
    seen_tc_ids = set()
    for row in rows:
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen_tc_ids:
                continue
            seen_tc_ids.add(cid)
            # Replace embedded newlines with space before lowercasing
            t = cell.text.replace("\n", " ").replace("\r", " ").strip().lower()
            if t:
                texts.append(t)
    return " ".join(texts)


def _match_fingerprint(header_text: str, is_library: bool) -> str:
    """
    Match header text against known fingerprints.

    ORDER MATTERS: More specific patterns must be checked before generic ones.
    """
    # Library attendance
    if is_library and "on roll" in header_text and "present" in header_text:
        return "library_attendance"

    # Library transactions
    if "particulars" in header_text:
        return "library_transactions"

    # Standard attendance — must have On Rolls + Absent
    if "on rolls" in header_text and "absent" in header_text:
        if "overall performance" in header_text or "overall dqi" in header_text:
            return "overall_staff_attendance"
        if (
            "category" in header_text
            or "teaching" in header_text
            or "dept" in header_text
        ):
            return "attendance"

    # Student Attendance (overall)
    if "ce" in header_text and "eee" in header_text and "cse" in header_text:
        return "overall_student_attendance"

    # Infrastructure — v3: "reported on" (with space) now matches because we
    # normalised newlines in _get_header_text()
    if "description" in header_text and (
        "reported" in header_text or "problem" in header_text
    ):
        return "infrastructure"

    # MTP specific sections
    if "mtp:" in header_text or "mtp (" in header_text:
        return "mtp_narrative"
    if "batch pills" in header_text or "batch pli" in header_text:
        return "mtp_batch_pills"

    # Participation BEFORE Events
    # v3: broadened — accept any combo of Name+Event+Status OR Name+Event+Date
    if (
        "name" in header_text
        and "event" in header_text
        and (
            "status" in header_text
            or "delegate" in header_text
            or "speaker" in header_text
            or "date" in header_text
        )
    ):
        return "staff_participation"  # Will be disambiguated by counter

    # Events
    if "seminar" in header_text or "workshop" in header_text or "short term" in header_text:
        return "events"
    if "event" in header_text and (
        "duration" in header_text or "participants" in header_text
    ):
        return "events"

    # Staff changes
    if "designation" in header_text and (
        "joining" in header_text or "leaving" in header_text
    ):
        return "staff_changes"

    # Classwork adjustments
    if "adjusted" in header_text and (
        "subject" in header_text or "dept" in header_text
    ):
        return "classwork"

    # Department Student Attendance tables (must not leak into Other/Narratives)
    if ("rolls" in header_text or "roll" in header_text) and ("present" in header_text or "absent" in header_text):
        return "student_attendance_table"

    # Incidents
    if "brief" in header_text and (
        "statement" in header_text or "remarks" in header_text
    ):
        return "incidents"

    return "unknown"


# ── Deterministic extractors ─────────────────────────────────────────────────
# NONE of these functions call the LLM. Numbers come only from table cells.

def _extract_attendance(table, dept_code: str, dept_name: str) -> dict:
    """
    Extract attendance data deterministically.
    Returns None if the table does not have the expected structure.
    Attendance numbers are NEVER passed to the LLM.
    """
    rows = _read_table_rows(table)
    if len(rows) < 2:  # Need at least header + 1 data row
        return None

    header = [h.lower().strip() for h in rows[0]]

    # Find column indices
    on_rolls_idx = _find_col(header, ["on rolls", "onrolls", "on roll"])
    absent_idx = _find_col(header, ["absent"])
    category_idx = _find_col(header, ["category"])

    if on_rolls_idx is None or absent_idx is None:
        return None

    teaching_count = 0
    non_teaching_count = 0
    total_on_rolls = 0
    total_absent = 0

    for row in rows[1:]:
        if len(row) <= max(on_rolls_idx, absent_idx):
            continue

        on_rolls = _parse_int(row[on_rolls_idx])
        absent = _parse_int(row[absent_idx])

        if on_rolls is None:
            continue

        # Determine teaching vs non-teaching
        category = (
            row[category_idx].lower().strip()
            if category_idx is not None and category_idx < len(row)
            else ""
        )

        if "non" in category:
            non_teaching_count = on_rolls
        else:
            teaching_count = on_rolls

        total_on_rolls += on_rolls or 0
        total_absent += absent or 0

    if total_on_rolls == 0:
        return None  # No valid data rows

    present = total_on_rolls - total_absent
    percentage = round(present / total_on_rolls * 100, 1) if total_on_rolls > 0 else None

    return {
        "dept": dept_name,
        "teaching_count": teaching_count or None,
        "non_teaching_count": non_teaching_count or None,
        "on_rolls": total_on_rolls,
        "absent": total_absent,
        "present": present,
        "percentage": percentage,
    }


def _extract_library_attendance(table) -> dict:
    """Extract library staff attendance."""
    rows = _read_table_rows(table)
    if len(rows) < 2:
        return None

    header = [h.lower().strip() for h in rows[0]]
    data = rows[1]

    on_rolls_idx = _find_col(header, ["on roll", "on rolls"])
    absent_with_idx = _find_col(header, ["absent with leave", "absent with"])
    absent_without_idx = _find_col(header, ["absent without leave", "absent without"])
    present_idx = _find_col(header, ["present"])

    return {
        "on_rolls": _safe_get_int(data, on_rolls_idx),
        "absent_with_leave": _safe_get_int(data, absent_with_idx),
        "absent_without_leave": _safe_get_int(data, absent_without_idx),
        "present": _safe_get_int(data, present_idx),
    }


def _extract_infrastructure(table, dept_name: str) -> list[dict]:
    """Extract infrastructure issues deterministically."""
    rows = _read_table_rows(table)
    if len(rows) < 2:
        return []

    header = [h.lower().strip() for h in rows[0]]
    desc_idx = _find_col(header, ["description", "problem"])
    reported_idx = _find_col(header, ["reported"])
    completed_idx = _find_col(header, ["completed"])
    remarks_idx = _find_col(header, ["remarks"])

    issues = []
    for row in rows[1:]:
        desc = _safe_get(row, desc_idx, "").strip()
        if not desc or _is_empty(desc):
            continue

        completed = _safe_get(row, completed_idx, "").strip()
        # Skip resolved issues (completed date is filled)
        if completed and not _is_empty(completed):
            continue

        issues.append({
            "dept": dept_name,
            "description": desc,
            "reported_on": _safe_get(row, reported_idx, ""),
            "status": "pending",
            "remarks": _safe_get(row, remarks_idx, "") or None,
        })

    return issues


def _extract_staff_changes(table, dept_name: str) -> list[dict]:
    """Extract staff joined/left entries."""
    rows = _read_table_rows(table)
    if len(rows) < 2:
        return []

    header = [h.lower().strip() for h in rows[0]]
    name_idx = _find_col(header, ["name", "faculty"])
    desig_idx = _find_col(header, ["designation"])
    date_idx = _find_col(header, ["joining", "leaving", "date"])
    remarks_idx = _find_col(header, ["remarks"])

    changes = []
    for row in rows[1:]:
        name = _safe_get(row, name_idx, "").strip()
        if not name or _is_empty(name):
            continue

        date_val = _safe_get(row, date_idx, "")
        remarks = _safe_get(row, remarks_idx, "")

        # Determine joined or left
        change_type = "joined"
        combined = (date_val + " " + remarks).lower()
        if "left" in combined or "leaving" in combined or "resign" in combined:
            change_type = "left"

        changes.append({
            "name": name,
            "dept": dept_name,
            "designation": _safe_get(row, desig_idx, ""),
            "type": change_type,
            "date": date_val,
        })

    return changes


def _extract_incidents(table, dept_name: str) -> list[dict]:
    """Extract discipline incidents."""
    rows = _read_table_rows(table)
    if len(rows) < 2:
        return []

    header = [h.lower().strip() for h in rows[0]]
    name_idx = _find_col(header, ["name"])
    id_idx = _find_col(header, ["r. no", "id no", "roll"])
    brief_idx = _find_col(header, ["brief", "statement"])
    remarks_idx = _find_col(header, ["remarks"])

    incidents = []
    for row in rows[1:]:
        name = _safe_get(row, name_idx, "").strip()
        brief = _safe_get(row, brief_idx, "").strip()
        if (not name or _is_empty(name)) and (not brief or _is_empty(brief)):
            continue

        incidents.append({
            "dept": dept_name,
            "type": "student",
            "name": name,
            "id": _safe_get(row, id_idx, "") or None,
            "brief": brief,
            "remarks": _safe_get(row, remarks_idx, "") or None,
        })

    return incidents


def _extract_library_data(table) -> tuple[dict, dict]:
    """Extract library transactions and services from the combined table."""
    rows = _read_table_rows(table)
    if len(rows) < 2:
        return {}, {}

    header = [h.lower().strip() for h in rows[0]]
    val_idx = len(rows[0]) - 1  # Last column is always the value column

    particulars_idx = _find_col(header, ["particulars", "description"])
    if particulars_idx is None:
        particulars_idx = 1

    transactions = {}
    services = {}

    TXN_RULES = [
        (lambda l: "books issued" in l or "check out" in l, "books_issued"),
        (lambda l: "books returned" in l or "check in" in l, "books_returned"),
        (lambda l: "total" in l and "visitors" in l and "lirc" in l, "visitors_lirc"),
        (lambda l: ("evening" in l or "5.00 pm" in l or "5 pm" in l or "5.00pm" in l), "visitors_evening_5_to_8"),
        (lambda l: "digital" in l, "visitors_digital"),
        (lambda l: "show" in l and "tell" in l and "visitor" in l, "show_and_tell_visitors"),
        (lambda l: "cvpc" in l or "c.v.p.c" in l, "cvpc_visitors"),
    ]

    SVC_RULES = [
        (lambda l: "plagiarism" in l or "turnitin" in l, "plagiarism_checks"),
        (lambda l: "show" in l and "tell" in l and "visitor" not in l, "show_and_tell"),
        (lambda l: "patent" in l, "patent_searches"),
        (lambda l: "scopus" in l, "scopus_searches"),
        (lambda l: "grammarly" in l, "grammarly_usage"),
        (lambda l: "duplicate" in l, "duplicate_id_cards"),
    ]

    for row in rows[1:]:
        label = _safe_get(row, particulars_idx, "").strip().lower()
        value = _safe_get_int(row, val_idx)

        if not label or value is None:
            continue

        matched = False
        for rule_fn, key in TXN_RULES:
            if rule_fn(label):
                transactions[key] = value
                matched = True
                break

        if not matched:
            for rule_fn, key in SVC_RULES:
                if rule_fn(label):
                    services[key] = value
                    break

    return transactions, services


# ── Table → text for LLM narrative sections ──────────────────────────────────

def _table_to_text(label: str, table) -> str:
    """Convert a table to pipe-delimited text for LLM consumption ONLY if it contains filled data rows."""
    rows = _read_table_rows(table)
    if not rows:
        return ""

    EMPTY_VALS = {"", "nil", "none", "no", "-", "—", "n/a", "na", "--", "i", "ii", "iii", "iv"}
    HEADER_KEYWORDS = (
        "name of the event", "for whom", "participants", "resource person",
        "s.no", "s. no", "particulars", "designation", "roll no", "r. no", "id no",
        "industry visited", "location of the industry", "no. of students visited",
        "staff/student", "date of the visit", "date and duration", "status (delegate",
        "delegate /paper", "delegate / paper"
    )

    # Find rows that have actual filled content in columns 1+
    data_rows = []
    for row in rows:
        row_str = " ".join(row).lower()
        # Skip header rows
        if any(h in row_str for h in HEADER_KEYWORDS):
            continue
        # Check if cells beyond S.No have substantive text (>2 chars, not empty/nil)
        substantive = [c.strip() for c in row[1:] if c.strip().lower() not in EMPTY_VALS and len(c.strip()) > 2]
        if substantive:
            data_rows.append(row)

    if not data_rows:
        return ""

    # Include label + header rows + data rows
    lines = [f"[{label}]"]
    for row in rows[:2]:
        lines.append(" | ".join(row))
    for row in data_rows:
        lines.append(" | ".join(row))

    return "\n".join(lines)


# ── Utility functions ─────────────────────────────────────────────────────────

def _read_table_rows(table) -> list[list[str]]:
    """
    Read all rows from a table, handling merged cells.

    Deduplication is done by checking the underlying XML tc element identity,
    not by comparing text content. This prevents data loss when adjacent cells
    legitimately contain the same value.
    """
    rows = []
    for row in table.rows:
        cells = []
        seen_tc_ids = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc_ids:
                continue  # Skip merged cell (same underlying XML element)
            seen_tc_ids.add(tc_id)
            text = cell.text.strip().replace("\n", " ").replace("\r", "")
            cells.append(text)
        rows.append(cells)
    return rows


def _find_col(header: list[str], keywords: list[str]) -> Optional[int]:
    """Find column index by keyword matching."""
    for i, h in enumerate(header):
        for kw in keywords:
            if kw in h:
                return i
    return None


def _parse_int(val: str) -> Optional[int]:
    """Parse an integer from a cell value, handling common formats."""
    if not val:
        return None
    stripped = val.strip()
    # Treat '--', '-', '—' as 0 (common in library reports for zero values)
    if stripped in ('--', '-', '—', 'nil', 'Nil', 'NIL'):
        return 0
    # Remove whitespace and common non-numeric chars
    cleaned = re.sub(r'[^\d]', '', stripped.split(':')[0].split('.')[0].strip())
    if not cleaned:
        return None
    try:
        return int(cleaned)
    except (ValueError, TypeError):
        return None


def _safe_get(lst: list, idx: Optional[int], default: str = "") -> str:
    """Safely get a value from a list by index."""
    if idx is None or idx >= len(lst):
        return default
    return lst[idx]


def _safe_get_int(lst: list, idx: Optional[int]) -> Optional[int]:
    """Safely get an integer value from a list by index."""
    val = _safe_get(lst, idx, "")
    return _parse_int(val)


def _is_empty(text: str) -> bool:
    """Check if a cell value is effectively empty."""
    if not text:
        return True
    normalized = text.strip().lower()
    return normalized in ("", "nil", "none", "no", "-", "—", "n/a", "na", "--")


def _has_real_content(text: str) -> bool:
    """Check if text block has meaningful narrative content beyond just headers."""
    lines = text.strip().split("\n")
    content_lines = [l for l in lines if not l.startswith("[") and l.strip()]
    if len(content_lines) <= 1:
        # A table with only 1 line is just a header with no data rows
        return False

    EMPTY_MARKERS = {"", "nil", "none", "no", "-", "—", "n/a", "na", "--", "rolls", "present", "absent", "i", "ii", "iii", "iv"}
    for line in content_lines[1:]:
        cells = [c.strip().lower() for c in line.split("|")]
        non_empty = [c for c in cells if c and c not in EMPTY_MARKERS and not c.isdigit() and len(c) > 2]
        if non_empty:
            return True

    return False


def _count_non_empty_rows(table) -> int:
    """Count data rows that have actual content (for classwork adjustments)."""
    rows = _read_table_rows(table)
    count = 0
    for row in rows[1:]:  # Skip header
        # A row has content if any cell (beyond S.No) has non-empty text
        if len(row) > 1 and any(not _is_empty(c) for c in row[1:]):
            count += 1
    return count
