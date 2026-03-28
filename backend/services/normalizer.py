import io
import os
import re
import zipfile


DEPT_TABLE_LABELS = [
    "Staff Attendance",
    "Infrastructure Issues or Maintenance",
    "Events / Seminars / Workshops",
    "Participation by Staff",
    "Participation by Students",
    "Staff Joined or Left",
    "Classwork Adjustments / Lecture Interchange",
    "Incidents (Discipline)",
    "Any Other Matter",
]

LIBRARY_TABLE_LABELS = [
    "Staff Attendance",
    "Infrastructure Issues or Maintenance",
    "Staff Joined or Left",
    "Library Services and Transactions",
    "Plagiarism / Show and Tell / Patent / Scopus / Grammarly / Duplicate IDs",
]


def normalize(file_bytes: bytes, filename: str, is_library: bool = False) -> str:
    """
    Extract text from a .docx or .pdf file.
    Returns structured text with labeled sections.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return _extract_docx(file_bytes, is_library)
    elif ext == ".pdf":
        return _extract_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .docx and .pdf are accepted.")


def truncate(text: str, max_chars: int = 15000) -> str:
    """Truncate text to max_chars. Gemini handles long context well, so default is generous."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n[...truncated for length]"


def extract_images(file_bytes: bytes, filename: str, dept_code: str) -> list[dict]:
    """
    Extract event-related images from a .docx file.

    Filters out:
    - The first image in every file (always the college logo/header)
    - Vector formats (EMF/WMF)
    - Very small images (< 5 KB — icons, bullets, decorations)

    Returns list of dicts: {dept_code, filename, image_bytes, content_type}
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext != ".docx":
        return []

    images = []
    try:
        zf = zipfile.ZipFile(io.BytesIO(file_bytes))
        media_files = sorted(
            [n for n in zf.namelist() if n.startswith("word/media/")]
        )

        for idx, media_path in enumerate(media_files):
            # Skip the FIRST image — it is always the college logo/header
            if idx == 0:
                continue

            try:
                img_bytes = zf.read(media_path)
                img_name = os.path.basename(media_path)
                ext_lower = os.path.splitext(img_name)[1].lower()

                # Skip vector formats (not real photos)
                if ext_lower in (".emf", ".wmf"):
                    continue

                # Skip small images (< 5 KB — icons, bullets, template decorations)
                if len(img_bytes) < 5120:
                    continue

                content_type_map = {
                    ".png": "image/png",
                    ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg",
                    ".gif": "image/gif",
                    ".bmp": "image/bmp",
                    ".tiff": "image/tiff",
                }
                content_type = content_type_map.get(ext_lower, "image/png")

                images.append({
                    "dept_code": dept_code,
                    "filename": f"{dept_code}_{img_name}",
                    "image_bytes": img_bytes,
                    "content_type": content_type,
                })
            except (zipfile.BadZipFile, Exception) as e:
                print(f"[WARNING] Failed to extract image {media_path} from {filename}: {e}")
                continue
    except Exception as e:
        print(f"[WARNING] Could not read zip for images from {filename}: {e}")

    return images


def _extract_docx(file_bytes: bytes, is_library: bool) -> str:
    """Extract structured text from a .docx file by reading its tables."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    labels = LIBRARY_TABLE_LABELS if is_library else DEPT_TABLE_LABELS
    output = []

    for i, table in enumerate(doc.tables):
        label = labels[i] if i < len(labels) else f"Section {i + 1}"
        rows = []

        for row in table.rows:
            cells = []
            prev_text = None
            for cell in row.cells:
                # Handle merged cells: skip duplicate cell text from merges
                text = cell.text.strip().replace("\n", " ").replace("\r", "")
                if text == prev_text:
                    continue  # Skip merged cell duplicates
                prev_text = text
                cells.append(text)

            # Skip completely empty rows
            if not any(c.strip() for c in cells):
                continue

            rows.append(" | ".join(cells))

        if rows:
            output.append(f"[{label}]")
            output.extend(rows)
            output.append("")

    # Also capture any loose paragraphs (free text outside tables)
    loose_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 5:
            loose_paras.append(text)

    if loose_paras:
        output.append("[Free Text / Other Notes]")
        output.extend(loose_paras)

    return "\n".join(output) if output else "[No content extracted]"


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")

    return "\n\n".join(pages) if pages else "[No content extracted]"
