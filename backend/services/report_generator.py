"""
DOCX Report Generator (Executive Edition)
=========================================
Generates a consolidated institutional daily report in .docx format matching
the VNRVJIET standard template with executive typography, clean table grids,
and formal layout.
"""

import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Color Palette (Executive Navy & Teal) ─────────────────────────────────────
DARK_BLUE = RGBColor(0x00, 0x20, 0x45)       # Primary Executive Navy
MID_BLUE = RGBColor(0x1E, 0x40, 0xAF)        # Secondary Deep Slate Blue
ACCENT_TEAL = RGBColor(0x0D, 0x94, 0x88)     # Accent Teal for section badges/bars
ACCENT_RED = RGBColor(0xB9, 0x1C, 0x1C)      # Missing / Incident Red
BODY_GRAY = RGBColor(0x1E, 0x29, 0x3B)       # Slate 800 body text
MUTED_GRAY = RGBColor(0x64, 0x74, 0x8B)      # Slate 500 metadata text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DARK_BLUE_HEX = "002045"
MID_BLUE_HEX = "1E40AF"
ACCENT_TEAL_HEX = "0D9488"
LIGHT_GRAY_HEX = "F8FAFC"                    # Subtle Slate 50 alternating rows
BORDER_HEX = "CBD5E1"                        # Subtle Slate 300 border


# ── XML & Styling Helpers ─────────────────────────────────────────────────────

def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set internal cell padding in dxa (twips, 1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold: bool = False, font_size: float = 9.0,
                   color: RGBColor = None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set text in a table cell with typography and line spacing."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = BODY_GRAY
    run.font.name = "Calibri"
    _set_cell_margins(cell)


def _add_header_row(table, headers: list[str]):
    """Style the first row of a table as an executive header."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        _set_cell_shading(cell, DARK_BLUE_HEX)
        _set_cell_text(cell, header, bold=True, font_size=9,
                       color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _add_data_row(table, values: list[str], row_idx: int, align_left_col: int = 1):
    """Add a data row with alternating slate/white background and border."""
    row = table.rows[row_idx]
    bg = LIGHT_GRAY_HEX if row_idx % 2 == 0 else "FFFFFF"
    for i, val in enumerate(values):
        cell = row.cells[i]
        _set_cell_shading(cell, bg)
        align = WD_ALIGN_PARAGRAPH.LEFT if i == align_left_col else WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_text(cell, str(val), font_size=9, alignment=align)


def _add_section_heading(doc, number: int, title: str):
    """Add a numbered section heading with navy bottom border."""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(f"{number}.  {title}")
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    pPr = heading._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="2" w:color="{DARK_BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_sub_heading(doc, title: str):
    """Add a sub-section heading in deep slate blue."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MID_BLUE
    run.font.name = "Calibri"


def _add_dept_heading(doc, dept_name: str):
    """Add a department name heading with teal accent bar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.2)

    # Accent bar
    run = p.add_run("▎ ")
    run.font.size = Pt(10.5)
    run.font.color.rgb = ACCENT_TEAL
    run.font.name = "Calibri"

    # Department name
    run = p.add_run(dept_name.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"


def _add_body_text(doc, text: str, bold_prefix: str = None, italic: bool = False, color: RGBColor = None):
    """Add a body paragraph with optional bold prefix and clean indentation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)

    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Calibri"

    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = color or BODY_GRAY
    run.font.name = "Calibri"
    run.italic = italic
    return p


def _insert_dept_images(doc, dept_code: str, all_images: list[dict], max_images: int = 3):
    """Insert event-related images for a specific department inline."""
    dept_images = [img for img in all_images if img["dept_code"] == dept_code]
    if not dept_images:
        return 0

    inserted = 0
    for img in dept_images[:max_images]:
        try:
            img_stream = io.BytesIO(img["image_bytes"])
            doc.add_picture(img_stream, width=Inches(3.8))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_p.paragraph_format.space_before = Pt(4)
            last_p.paragraph_format.space_after = Pt(4)
            inserted += 1
        except Exception as e:
            print(f"[WARNING] Failed to insert image {img['filename']}: {e}")

    return inserted


def _val(v):
    """Format a value for display — use '—' for None/null."""
    if v is None or v == "":
        return "—"
    return str(v)


def _pct(val):
    """Format percentage."""
    if val is None:
        return "—"
    return f"{val:.1f}%"


# ── Section 1: Staff Attendance Summary Only ──────────────────────────────────

def _build_staff_attendance(doc, report: dict, section_num: int) -> int:
    """
    Build Section 1: Staff Attendance Summary.
    Uses attendance.departments (deterministic) directly for a clean department-wise table.
    Excludes Library staff attendance and student attendance tables.
    """
    depts = report.get("attendance", {}).get("departments", [])
    if not depts:
        return section_num

    _add_section_heading(doc, section_num, "Staff Attendance Summary")

    headers = [
        "S.No", "Department", "Teaching", "Non-Teaching",
        "On Rolls", "Absent", "Present", "Attendance %"
    ]
    table = doc.add_table(rows=1 + len(depts) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _add_header_row(table, headers)

    total_teach = total_nt = total_rolls = total_absent = total_present = 0
    for i, d in enumerate(depts):
        t  = d.get("teaching_count", 0) or 0
        nt = d.get("non_teaching_count", 0) or 0
        on = d.get("on_rolls", 0) or 0
        ab = d.get("absent", 0) or 0
        pr = d.get("present", 0) or 0
        total_teach   += t
        total_nt      += nt
        total_rolls   += on
        total_absent  += ab
        total_present += pr

        _add_data_row(table, [
            str(i + 1), d.get("dept", ""),
            _val(t) if t else "—", _val(nt) if nt else "—",
            _val(on), _val(ab), _val(pr),
            _pct(d.get("percentage")),
        ], i + 1, align_left_col=1)

    # Totals Row
    total_pct = round(total_present / total_rolls * 100, 1) if total_rolls else 0
    total_row = table.rows[-1]
    totals = [
        "", "TOTAL",
        str(total_teach) if total_teach else "—",
        str(total_nt)    if total_nt    else "—",
        str(total_rolls), str(total_absent), str(total_present), _pct(total_pct)
    ]
    for i, val in enumerate(totals):
        cell = total_row.cells[i]
        _set_cell_shading(cell, DARK_BLUE_HEX)
        _set_cell_text(cell, val, bold=True, font_size=9,
                       color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return section_num + 1


# ── Section 2: Department Highlights (Events & Key Activities) ────────────────

def _build_department_highlights(doc, report: dict, section_num: int, all_images: list[dict]) -> int:
    """
    Build Section 2: Department Highlights.
    Covers all standard academic departments:
    - Active: lists polished event bullets + other matters + inline images.
    - No highlights: 'No significant highlights reported for today.'
    - Missing report: 'Report not submitted / data not available for today.'
    """
    highlights = report.get("department_highlights", [])
    if not highlights:
        return section_num

    _add_section_heading(doc, section_num, "Department Highlights — Events & Key Activities")

    for dept_block in highlights:
        dept_name = dept_block.get("dept", "Unknown Department")
        dept_code = dept_block.get("dept_code", "").lower()
        status = dept_block.get("status", "active")
        events = dept_block.get("events", [])
        other_matters = dept_block.get("other_matters", [])

        # MTP and Library have their own dedicated sections
        if dept_code in ("mtp", "library"):
            continue

        _add_dept_heading(doc, dept_name)

        if status == "missing_report":
            _add_body_text(doc, "Report not submitted / data not available for today.",
                           italic=True, color=ACCENT_RED)
            continue

        if status == "no_highlights" or (not events and not other_matters):
            _add_body_text(doc, "No significant highlights reported for today.",
                           italic=True, color=MUTED_GRAY)
            continue

        # Render Active Events
        for ev in events:
            name = ev.get("name", "Department Activity")
            summary = ev.get("summary", "")
            date_str = ev.get("date", "")
            duration = ev.get("duration", "")
            internal = ev.get("participants_internal")
            external = ev.get("participants_external")
            resource_person = ev.get("resource_person")

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)

            run = p.add_run(f"● {name}")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = MID_BLUE
            run.font.name = "Calibri"

            meta_parts: list[str] = []
            if date_str:
                meta_parts.append(f"Date: {date_str}")
            if duration:
                meta_parts.append(f"Duration: {duration}")
            if resource_person:
                meta_parts.append(f"Speaker: {resource_person}")
            if internal is not None:
                meta_parts.append(f"Participants: {internal}")
            elif external is not None:
                meta_parts.append(f"Participants: {external}")

            if meta_parts:
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_before = Pt(0)
                meta_p.paragraph_format.space_after = Pt(2)
                meta_p.paragraph_format.left_indent = Cm(1.0)
                run_m = meta_p.add_run("  •  ".join(meta_parts))
                run_m.font.size = Pt(8.5)
                run_m.font.color.rgb = MUTED_GRAY
                run_m.font.name = "Calibri"
                run_m.italic = True

            if summary:
                _add_body_text(doc, summary)

        # Other Matters
        if other_matters:
            for matter in other_matters:
                if isinstance(matter, str) and matter.strip():
                    _add_body_text(doc, matter.strip(), bold_prefix="• Note: ")
                elif isinstance(matter, dict):
                    desc = matter.get("description", "")
                    if desc.strip():
                        _add_body_text(doc, desc.strip(), bold_prefix="• Note: ")

        # Insert images inline
        img_count = _insert_dept_images(doc, dept_code, all_images, max_images=3)
        if img_count > 0:
            print(f"[INFO] Inserted {img_count} image(s) for {dept_code.upper()}")

    return section_num + 1


# ── Section 3: Mentoring, Training & Placements (MTP) ─────────────────────────

def _parse_batch_pills_table(text: str):
    """
    Extract the header + values pipe-separated table from batch pills text.
    Returns (headers: list[str], values: list[str]) or (None, None).
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    BRANCH_KEYS = {"CSE", "AIML", "IOT", "IT", "ECE", "EEE", "EIE", "ME", "CE", "AE"}
    header_idx = None
    for i, line in enumerate(lines):
        cols = [c.strip() for c in line.split("|") if c.strip()]
        if len(cols) >= 8 and len(BRANCH_KEYS & set(cols)) >= 4:
            header_idx = i
            break
    if header_idx is None:
        return None, None
    headers = [c.strip() for c in lines[header_idx].split("|") if c.strip()]
    if header_idx + 1 < len(lines):
        val_line = lines[header_idx + 1]
        vals = [c.strip() for c in val_line.split("|") if c.strip()]
        if all(v.isdigit() for v in vals):
            return headers, vals
    return headers, None


def _build_mtp_sections(doc, report: dict, section_num: int) -> int:
    """
    Build Section 3: Mentoring, Training & Placements (MTP).
    Preserves source order of placement drives/trainings and formats Batch Pills table.
    """
    mtp_narrative = report.get("mtp_narrative", "").strip()
    batch_pills   = report.get("mtp_batch_pills", "").strip()
    mtp_summary   = report.get("mtp_summary", [])

    if not mtp_narrative and not batch_pills and not mtp_summary:
        return section_num

    _add_section_heading(doc, section_num, "Mentoring, Training & Placements (MTP)")

    # ── MTP Structured Activity Items (In Source Order) ──────────────────────
    if mtp_summary:
        _add_sub_heading(doc, "Placement & Training Activity Highlights")

        ACTIVITY_LABELS = {
            "placement_drive": "Placement Drive",
            "ppt":             "Pre-Placement Talk",
            "aptitude_test":   "Aptitude Assessment",
            "training":        "Training Session",
            "mock_interview":  "Mock Interview",
            "internship":      "Internship",
            "other":           "Activity",
        }

        for item in mtp_summary:
            if not isinstance(item, dict):
                continue
            company       = item.get("company") or ""
            activity_type = item.get("activity_type", "other")
            summary_text  = item.get("summary", "").strip()
            student_count = item.get("student_count")
            batch         = item.get("batch") or ""
            status        = item.get("status") or ""

            if not summary_text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Cm(0.5)

            label = company if company else ACTIVITY_LABELS.get(activity_type, "Placement Activity")
            run = p.add_run(f"● {label}")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"

            badge_label = ACTIVITY_LABELS.get(activity_type, activity_type.replace("_", " ").title())
            run2 = p.add_run(f"  [{badge_label}]")
            run2.bold = True
            run2.font.size = Pt(8.5)
            run2.font.color.rgb = ACCENT_TEAL
            run2.font.name = "Calibri"

            meta_parts = []
            if batch:
                meta_parts.append(batch)
            if status:
                meta_parts.append(status)
            if student_count is not None:
                meta_parts.append(f"{student_count} Students")

            if meta_parts:
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_before = Pt(0)
                meta_p.paragraph_format.space_after  = Pt(2)
                meta_p.paragraph_format.left_indent  = Cm(1.0)
                run_m = meta_p.add_run("  •  ".join(meta_parts))
                run_m.font.size = Pt(8.5)
                run_m.font.color.rgb = MUTED_GRAY
                run_m.font.name = "Calibri"
                run_m.italic = True

            _add_body_text(doc, summary_text)

    elif mtp_narrative:
        _add_sub_heading(doc, "Placement & Training Activity Highlights")
        text = mtp_narrative.replace("[MTP Section IV]", "").strip()
        for line in text.split("\n"):
            line = line.strip()
            if line:
                _add_body_text(doc, line)

    # ── Batch Pills Table (As is) ─────────────────────────────────────────────
    if batch_pills:
        _add_sub_heading(doc, "Batch Pills Open Summary")
        text = batch_pills.replace("[Batch Pills Open Summary]", "").strip()
        headers, values = _parse_batch_pills_table(text)
        if headers and values:
            n_cols = min(len(headers), len(values))
            tbl = doc.add_table(rows=2, cols=n_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = 'Table Grid'
            for i, h in enumerate(headers[:n_cols]):
                cell = tbl.rows[0].cells[i]
                _set_cell_shading(cell, DARK_BLUE_HEX)
                _set_cell_text(cell, h, bold=True, font_size=8,
                               color=WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            for i, v in enumerate(values[:n_cols]):
                cell = tbl.rows[1].cells[i]
                _set_cell_shading(cell, LIGHT_GRAY_HEX)
                _set_cell_text(cell, v, bold=True, font_size=9,
                               color=DARK_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            for line in text.split("\n"):
                line = line.strip()
                if line and not line.startswith("["):
                    _add_body_text(doc, line)

    return section_num + 1


# ── Section 4: Faculty & Student Participation ────────────────────────────────

def _build_participation(doc, report: dict, section_num: int) -> int:
    """
    Build Section 4: External Participation.
    Separated into 4.1 Faculty & Staff and 4.2 Student Participation,
    ordered by importance without raw tags.
    """
    staff_p = report.get("staff_participation", [])
    student_p = report.get("student_participation", [])

    if not staff_p and not student_p:
        return section_num

    _add_section_heading(doc, section_num, "Participation by Faculty & Students (External)")

    # 4.1 Faculty & Staff Participation
    if staff_p:
        _add_sub_heading(doc, "4.1 Faculty & Staff Development / Participation")
        for s in staff_p:
            name = s.get("name", "")
            dept = s.get("dept", "")
            event = s.get("event", "")
            role = s.get("role", s.get("status", ""))
            date_str = s.get("date", "")
            venue = s.get("venue", "")
            summary = s.get("summary", "")

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)

            run = p.add_run(f"● {name}")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"

            detail_parts: list[str] = []
            if dept:
                detail_parts.append(dept.upper())
            if role:
                detail_parts.append(role)
            if event:
                detail_parts.append(f'at "{event}"')
            if venue:
                detail_parts.append(f"({venue})")
            if date_str:
                detail_parts.append(f"on {date_str}")

            if detail_parts:
                run = p.add_run(f"  — {', '.join(detail_parts)}")
                run.font.size = Pt(9)
                run.font.color.rgb = BODY_GRAY
                run.font.name = "Calibri"

            if summary:
                _add_body_text(doc, summary)

    # 4.2 Student Achievements & Participation
    if student_p:
        _add_sub_heading(doc, "4.2 Student Achievements & Participation")
        for s in student_p:
            name = s.get("name", "")
            dept = s.get("dept", "")
            event = s.get("event", "")
            achievement = s.get("achievement", s.get("status", ""))
            date_str = s.get("date", "")
            summary = s.get("summary", "")

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)

            run = p.add_run(f"● {name}")
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_BLUE
            run.font.name = "Calibri"

            detail_parts: list[str] = []
            if dept:
                detail_parts.append(dept.upper())
            if event:
                detail_parts.append(f'at "{event}"')
            if achievement:
                detail_parts.append(f"— {achievement}")
            if date_str:
                detail_parts.append(f"({date_str})")

            if detail_parts:
                run = p.add_run(f"  — {', '.join(detail_parts)}")
                run.font.size = Pt(9)
                run.font.color.rgb = BODY_GRAY
                run.font.name = "Calibri"

            if summary:
                _add_body_text(doc, summary)

    return section_num + 1


# ── Operational Sections (5 to 9) ────────────────────────────────────────────

def _build_classwork_adjustments(doc, report: dict, section_num: int) -> int:
    """Build Section 5: Classwork Adjustments."""
    adjustments = report.get("classwork_adjustments", [])
    if not adjustments:
        return section_num

    _add_section_heading(doc, section_num, "Classwork Adjustments / Lecture Interchange")

    headers = ["S.No", "Department", "Number of Adjustments"]
    table = doc.add_table(rows=1 + len(adjustments), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _add_header_row(table, headers)

    for i, adj in enumerate(adjustments):
        _add_data_row(table, [
            str(i + 1), adj.get("dept", "").upper(), _val(adj.get("count")),
        ], i + 1, align_left_col=1)

    return section_num + 1


def _build_staff_changes(doc, report: dict, section_num: int) -> int:
    """Build Section 6: Staff Movement (Joined / Left)."""
    changes = report.get("staff_changes", [])
    if not changes:
        return section_num

    _add_section_heading(doc, section_num, "Staff Movement (Joined / Relieved)")

    headers = ["S.No", "Name", "Department", "Designation", "Status", "Date"]
    table = doc.add_table(rows=1 + len(changes), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _add_header_row(table, headers)

    for i, c in enumerate(changes):
        _add_data_row(table, [
            str(i + 1), c.get("name", "—"), c.get("dept", "—").upper(),
            c.get("designation", "—"), (c.get("type") or "Joined").upper(), c.get("date", "—"),
        ], i + 1, align_left_col=1)

    return section_num + 1


def _build_infrastructure(doc, report: dict, section_num: int) -> int:
    """Build Section 7: Infrastructure Issues / Maintenance (Pending)."""
    issues = [
        i for i in report.get("infrastructure_issues", [])
        if str(i.get("status") or "pending").lower() == "pending"
    ]
    if not issues:
        return section_num

    _add_section_heading(doc, section_num, "Infrastructure Issues / Maintenance (Pending)")

    headers = ["S.No", "Department", "Description", "Reported On", "Remarks"]
    table = doc.add_table(rows=1 + len(issues), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _add_header_row(table, headers)

    for i, iss in enumerate(issues):
        _add_data_row(table, [
            str(i + 1), iss.get("dept", "").upper(),
            iss.get("description", "—"), _val(iss.get("reported_on")),
            _val(iss.get("remarks")),
        ], i + 1, align_left_col=2)

    return section_num + 1


def _build_incidents(doc, report: dict, section_num: int) -> int:
    """Build Section 8: Disciplinary Incidents."""
    incidents = report.get("incidents", [])
    if not incidents:
        return section_num

    _add_section_heading(doc, section_num, "Incidents / Disciplinary Actions")

    headers = ["S.No", "Department", "Type", "Name", "Brief Description", "Remarks"]
    table = doc.add_table(rows=1 + len(incidents), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _add_header_row(table, headers)

    for i, inc in enumerate(incidents):
        _add_data_row(table, [
            str(i + 1), inc.get("dept", "").upper(), inc.get("type", "—"),
            inc.get("name", "—"), inc.get("brief") or inc.get("description") or "—", _val(inc.get("remarks")),
        ], i + 1, align_left_col=4)

    return section_num + 1


def _build_library(doc, report: dict, section_num: int) -> int:
    """Build Section 9: Library Services & Transactions."""
    txn = report.get("library_transactions", {})
    svc = report.get("library_services", {})

    if not txn and not svc:
        return section_num

    _add_section_heading(doc, section_num, "Library Services & Transactions")

    if txn:
        _add_sub_heading(doc, "Library Transactions & Footfall")
        txn_items = [
            ("Books Issued (Check Out)", txn.get("books_issued")),
            ("Books Returned (Check In)", txn.get("books_returned")),
            ("Today's Visitors to LIRC", txn.get("visitors_lirc")),
            ("Evening Users (5 PM – 8 PM)", txn.get("visitors_evening_5_to_8")),
            ("Digital Library Visitors", txn.get("visitors_digital")),
            ("Show & Tell Visitors", txn.get("show_and_tell_visitors")),
            ("CVPC Visitors", txn.get("cvpc_visitors")),
        ]
        txn_items = [(k, v) for k, v in txn_items if v is not None]
        if txn_items:
            table = doc.add_table(rows=1 + len(txn_items), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            _add_header_row(table, ["Particulars", "Count"])
            for i, (label, value) in enumerate(txn_items):
                _add_data_row(table, [label, _val(value)], i + 1, align_left_col=0)

    if svc:
        has_svc = any(v for v in svc.values() if v is not None)
        if has_svc:
            _add_sub_heading(doc, "Digital & Research Services")
            svc_items = [
                ("Plagiarism Checks (Turnitin)", svc.get("plagiarism_checks")),
                ("Show & Tell", svc.get("show_and_tell")),
                ("Patent Searches", svc.get("patent_searches")),
                ("Scopus Indexing Service", svc.get("scopus_searches")),
                ("Grammarly Usage", svc.get("grammarly_usage")),
                ("Duplicate ID Cards Issued", svc.get("duplicate_id_cards")),
            ]
            svc_items = [(k, v) for k, v in svc_items if v is not None]
            if svc_items:
                table = doc.add_table(rows=1 + len(svc_items), cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                _add_header_row(table, ["Service", "Count"])
                for i, (label, value) in enumerate(svc_items):
                    _add_data_row(table, [label, _val(value)], i + 1, align_left_col=0)

    return section_num + 1


# ── Main Entry Point ──────────────────────────────────────────────────────────

def generate_docx(report: dict, output_path: str = None,
                  all_images: list[dict] = None) -> bytes:
    """
    Generate the formatted executive daily report DOCX from the consolidated JSON.
    """
    if all_images is None:
        all_images = []

    doc = Document()

    # ── Page setup (A4, 1.8cm margins) ────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Format date ───────────────────────────────────────────────────────────
    report_date_str = report.get("report_date", "—")
    try:
        dt = datetime.strptime(report_date_str, "%Y-%m-%d")
        formatted_date = dt.strftime("%d %B %Y")
    except Exception:
        formatted_date = report_date_str

    # ── Title Banner ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("VNR Vignana Jyothi Institute of Engineering & Technology")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(f"Consolidated Institutional Daily Report  |  {formatted_date}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED_GRAY
    run.font.name = "Calibri"

    # Accent Divider Rule
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(0)
    hr.paragraph_format.space_after = Pt(8)
    pPr = hr._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{DARK_BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Sequential Sections ───────────────────────────────────────────────────
    num = 1
    num = _build_staff_attendance(doc, report, num)
    num = _build_department_highlights(doc, report, num, all_images)
    num = _build_mtp_sections(doc, report, num)
    num = _build_participation(doc, report, num)
    num = _build_classwork_adjustments(doc, report, num)
    num = _build_staff_changes(doc, report, num)
    num = _build_infrastructure(doc, report, num)
    num = _build_incidents(doc, report, num)
    num = _build_library(doc, report, num)

    # ── Institutional Footer Note ─────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Consolidated by AI  |  VNRVJIET Principal's Office")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED_GRAY
    run.font.name = "Calibri"
    run.italic = True

    # ── Save buffer ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    if output_path:
        with open(output_path, "wb") as f:
            f.write(docx_bytes)

    return docx_bytes